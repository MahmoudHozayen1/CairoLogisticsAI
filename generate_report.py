"""Generate the graduation-project report (docs/SwiftRoute_Report.docx).

Run::

    python generate_report.py

The report is a complete, defense-ready document: title page, abstract, an
auto-updating table of contents, chapters covering the problem, background
theory (VRP / TSP / heuristics), the system architecture, the optimisation
methodology, the benchmark experiment and its results (with the charts and
numbers produced by ``scripts/benchmark_optimizer.py``), testing, and the
conclusion.

Live benchmark numbers are read from ``docs/benchmarks/summary.json`` when
present, so regenerating after a new benchmark run keeps the report in sync.
"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
BENCH_DIR = os.path.join(HERE, "docs", "benchmarks")
SUMMARY_PATH = os.path.join(BENCH_DIR, "summary.json")
OUT = os.path.join(HERE, "docs", "SwiftRoute_Report.docx")

# Brand palette (matches the web app + slide deck).
PRIMARY = RGBColor(0x0D, 0x3B, 0x66)
ACCENT = RGBColor(0xEE, 0x6C, 0x4D)
DARK = RGBColor(0x1F, 0x29, 0x33)

# Human-readable order + labels if the summary is unavailable.
FALLBACK_ORDER = ["fifo", "nearest", "two_opt", "or_opt", "christofides", "annealing"]
FALLBACK_LABELS = {
    "fifo": "As received (FIFO)",
    "nearest": "Nearest Neighbour",
    "two_opt": "Nearest Neighbour + 2-opt",
    "or_opt": "2-opt + Or-opt (thorough)",
    "christofides": "Christofides (NetworkX)",
    "annealing": "Simulated Annealing (NetworkX)",
}


# --------------------------------------------------------------------------- #
#  Low-level helpers
# --------------------------------------------------------------------------- #
def load_summary():
    try:
        with open(SUMMARY_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def add_toc(doc):
    """Insert a Word table-of-contents field (updates on right-click > Update Field)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose \u201cUpdate Field\u201d to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(node)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = PRIMARY
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = PRIMARY
    return p


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def image(doc, filename, width_in=6.1, caption=None):
    path = os.path.join(BENCH_DIR, filename)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK


# --------------------------------------------------------------------------- #
#  Content sections
# --------------------------------------------------------------------------- #
def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("SwiftRoute")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("An AI-Assisted Logistics Platform with a\nBenchmarked Route-Optimisation Engine")
    run.font.size = Pt(18)
    run.font.color.rgb = DARK

    for _ in range(6):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Graduation Project\nDepartment of Computer Science\n\nAuthor: ______________________\nSupervisor: ______________________\nDate: ______________________")
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    doc.add_page_break()


def abstract(doc):
    h1(doc, "Abstract")
    para(doc,
         "SwiftRoute is a complete, real-world logistics and last-mile delivery platform in which "
         "merchants create shipments, an optimisation engine plans courier routes, couriers deliver "
         "with photo proof, and customers track parcels in real time. Its technical core is a route "
         "optimiser that decomposes the NP-hard Vehicle Routing Problem into geographic clustering, "
         "capacity-aware courier assignment, and Travelling-Salesman sequencing.")
    para(doc,
         "This report focuses on the data-science contribution: six interchangeable sequencing "
         "techniques \u2014 from a naive FIFO baseline through Nearest-Neighbour, 2-opt and Or-opt local "
         "search to the graph-theoretic Christofides algorithm and a Simulated-Annealing metaheuristic "
         "(the latter two via NetworkX). A reproducible benchmark harness compares all six across "
         "hundreds of synthetic scenarios, measuring distance, improvement over the baseline, the gap "
         "to the exact optimum on small instances, and runtime. The optimised techniques reduce total "
         "courier distance by roughly 60% versus naive dispatch and come within a fraction of a percent "
         "of optimal, giving an empirically grounded justification for the engine's design.")
    para(doc,
         "The platform also carries an explainable predictive layer built on scikit-learn: "
         "cross-validated ETA and late-risk models, a transparent seasonal demand forecaster, "
         "handling-note NLP, a reinforcement-learned pointer policy for routing, and courier-behaviour "
         "clustering. Every model reports its reasoning and is validated with cross-validation against a "
         "naive baseline, so the accuracy figures reflect genuine generalisation rather than a favourable "
         "data split.")
    doc.add_page_break()


def toc_page(doc):
    h1(doc, "Contents")
    add_toc(doc)
    doc.add_page_break()


def introduction(doc):
    h1(doc, "1. Introduction")
    para(doc,
         "E-commerce growth has made last-mile delivery the most expensive and time-sensitive stage of "
         "the logistics chain. Manual courier dispatch is slow, uneven and error-prone, while customers "
         "increasingly expect real-time visibility of their parcels. SwiftRoute addresses both sides of "
         "this problem: an operational platform for the business and an optimisation engine that plans "
         "efficient courier routes automatically.")
    h2(doc, "1.1 Problem Statement")
    para(doc,
         "Given a set of parcels held at a hub and a fleet of couriers with limited vehicle capacity, "
         "assign parcels to couriers and order each courier's stops so that total travel distance (and "
         "therefore time and cost) is minimised, while respecting capacity and avoiding closed roads.")
    h2(doc, "1.2 Objectives")
    bullets(doc, [
        "Build a complete, database-agnostic web platform with role-based portals (admin, courier, merchant) and public tracking.",
        "Implement an optimisation engine that clusters parcels, assigns them under capacity limits, and sequences each route.",
        "Provide multiple, interchangeable optimisation techniques spanning heuristics, an approximation algorithm, and a metaheuristic.",
        "Empirically benchmark the techniques and quantify their savings, optimality gap and runtime.",
        "Add an explainable predictive layer (ETA, late-risk, demand forecasting, NLP, learning-to-route and behaviour personas), validated with cross-validation against naive baselines.",
        "Keep heavy scientific dependencies optional so the system always runs.",
    ])


def background(doc):
    h1(doc, "2. Background and Related Work")
    h2(doc, "2.1 The Vehicle Routing Problem")
    para(doc,
         "The Vehicle Routing Problem (VRP) generalises the Travelling Salesman Problem (TSP) to multiple "
         "vehicles with capacity constraints. Both are NP-hard: the number of possible routes grows "
         "factorially with the number of stops, so exact solutions are only tractable for very small "
         "instances. Practical systems therefore rely on heuristics and metaheuristics that find "
         "near-optimal routes quickly.")
    h2(doc, "2.2 Techniques Used in this Project")
    bullets(doc, [
        "k-means clustering \u2014 partitions drop-off points into geographically coherent zones, one per courier.",
        "Nearest-Neighbour \u2014 a greedy TSP heuristic that always moves to the closest unvisited stop.",
        "2-opt \u2014 local search that repeatedly reverses route segments to remove crossing or expensive edges.",
        "Or-opt \u2014 local search that relocates short chains of 1\u20133 stops to cheaper positions.",
        "Christofides' algorithm \u2014 a metric-TSP approximation with a proven 1.5\u00d7-optimal guarantee for the closed tour.",
        "Simulated Annealing \u2014 a metaheuristic that escapes local optima by occasionally accepting worse moves under a cooling schedule.",
    ])
    para(doc,
         "The last two are provided by NetworkX, a mature Python graph library. Because they are optional "
         "enhancements, the engine falls back to its pure-Python 2-opt sequence when NetworkX is absent, so "
         "dispatch never fails.")


def architecture(doc):
    h1(doc, "3. System Architecture")
    para(doc,
         "SwiftRoute is a single Flask application built with the application-factory pattern and organised "
         "into blueprints for each concern: authentication, admin operations, the courier portal, the "
         "merchant portal, public tracking and a JSON API. Persistence uses SQLAlchemy and is "
         "database-agnostic \u2014 SQLite by default for zero-configuration local use, and PostgreSQL in "
         "production via a single environment variable.")
    h2(doc, "3.1 Core Data Model")
    bullets(doc, [
        "User \u2014 admins, couriers and merchants in one table, separated by role.",
        "Hub \u2014 a warehouse with a fleet of couriers and geographic coordinates.",
        "Shipment \u2014 a parcel with a full lifecycle, receiver location and cash-on-delivery amount.",
        "ShipmentEvent \u2014 an append-only timeline; RouteStop \u2014 persisted route geometry and ETA.",
        "RoadClosure \u2014 an admin-managed area the optimiser routes around.",
    ])
    h2(doc, "3.2 Front End and Routing Geometry")
    para(doc,
         "The interface is server-rendered with Jinja2 and Bootstrap, using Leaflet for interactive maps "
         "and Chart.js for dashboards. Routes are drawn along real streets via the OSRM service (results "
         "cached on disk) with a straight-line fallback when offline, and a time-of-day traffic simulation "
         "colours each segment by congestion.")


def methodology(doc):
    h1(doc, "4. Optimisation Methodology")
    para(doc, "For each hub, the engine runs a four-stage pipeline:")
    numbered(doc, [
        "Cluster: k-means partitions at-warehouse parcels into one zone per available courier.",
        "Assign: a capacity-aware greedy pass gives each parcel to its nearest courier that still has room; parcels beyond total fleet capacity are held for the next round.",
        "Sequence: the chosen technique orders each courier's stops (the TSP sub-problem).",
        "Geometry and ETA: the route is drawn along real roads, avoiding active closures, and an ETA per stop is derived from distance, average speed and the simulated traffic at the planned dispatch time.",
    ])
    h2(doc, "4.1 Open Routes vs. Closed Tours")
    para(doc,
         "Courier routes are open: a courier starts at the hub and finishes at the last delivery without "
         "returning. Christofides, however, optimises a closed cycle. The implementation therefore solves "
         "the cycle and then cuts it at the depot in whichever of the two orientations yields the shorter "
         "open route \u2014 an honest adaptation whose consequences are visible in the results.")


def experiment(doc, summary):
    h1(doc, "5. Experimental Setup")
    para(doc,
         "The benchmark harness (scripts/benchmark_optimizer.py) evaluates every technique on synthetic "
         "delivery scenarios generated around the default service centre, with a fixed random seed for full "
         "reproducibility. It requires no database or network.")
    sizes = summary.get("sizes", [8, 15, 25, 40]) if summary else [8, 15, 25, 40]
    rows = [
        ("Random seed", "42"),
        ("Scenarios per size", "40"),
        ("Stops per route", ", ".join(str(s) for s in sizes)),
        ("Exact optimum", "Brute force for routes with \u2264 8 stops"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    para(doc, "")
    para(doc, "Four metrics are recorded per route:")
    bullets(doc, [
        "Distance (km) \u2014 total open-route distance from the depot through all stops.",
        "Improvement vs. FIFO (%) \u2014 distance saved relative to creation-order dispatch.",
        "Optimality gap (%) \u2014 distance above the exact optimum, where brute force is feasible.",
        "Runtime (ms) \u2014 wall-clock sequencing time.",
    ])


def results(doc, summary):
    h1(doc, "6. Results and Discussion")
    order = summary.get("order", FALLBACK_ORDER) if summary else FALLBACK_ORDER
    overall = summary.get("overall", {}) if summary else {}

    h2(doc, "6.1 Aggregate Comparison")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(["Technique", "Avg distance (km)", "Improvement vs FIFO", "Optimality gap", "Runtime (ms)"]):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for key in order:
        s = overall.get(key, {})
        label = s.get("label", FALLBACK_LABELS.get(key, key))
        dist = s.get("distance_km")
        imp = s.get("improvement_pct")
        gap = s.get("optimality_gap_pct")
        rt = s.get("runtime_ms")
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = f"{dist:.2f}" if dist is not None else "\u2013"
        cells[2].text = f"{imp:.1f}%" if imp is not None else "\u2013"
        cells[3].text = f"{gap:.1f}%" if gap is not None else "n/a"
        cells[4].text = f"{rt:.2f}" if rt is not None else "\u2013"
    para(doc, "")

    h2(doc, "6.2 Distance and Savings")
    image(doc, "distance_by_strategy.png",
          caption="Figure 1. Average total route distance by technique (lower is better).")
    image(doc, "improvement_by_strategy.png",
          caption="Figure 2. Average distance saved versus the FIFO baseline (higher is better).")
    para(doc,
         "Every optimisation technique beats naive FIFO dispatch by a wide margin \u2014 the strongest "
         "methods save roughly 60\u201362% of total distance on average. This is the headline justification "
         "for automated route planning.")

    h2(doc, "6.3 Optimality")
    image(doc, "optimality_gap.png",
          caption="Figure 3. Gap above the exact optimum on 8-stop instances (lower is better).")
    para(doc,
         "On small instances where the exact optimum is computable, Or-opt sits about 0.1% above optimal, "
         "and both 2-opt and Simulated Annealing are below 1%. Nearest-Neighbour and Christofides trail; "
         "Christofides' larger gap reflects that its guarantee is for a closed tour while couriers run open "
         "routes.")

    h2(doc, "6.4 Runtime and Scaling")
    image(doc, "runtime_by_strategy.png",
          caption="Figure 4. Average runtime per route by technique.")
    image(doc, "runtime_scaling.png",
          caption="Figure 5. Runtime growth as the number of stops increases (log scale).")
    image(doc, "distance_distribution.png",
          caption="Figure 6. Distribution of route distance at the largest instance size.")
    para(doc,
         "The runtime view exposes the quality/speed trade-off. Nearest-Neighbour is effectively instant "
         "but leaves distance on the table; Or-opt is the most thorough but grows fastest; 2-opt offers the "
         "best balance and is the production default. Christofides scales gently, making it an attractive "
         "graph-based option for larger routes.")


def ml_layer(doc):
    h1(doc, "7. Predictive Intelligence Layer (AI/ML)")
    para(doc,
         "Beyond route optimisation, SwiftRoute includes an explainable machine-learning layer that turns "
         "the platform into a genuine data-science system. It is built entirely on numpy, scikit-learn and "
         "pandas \u2014 no deep-learning runtime \u2014 so it trains in seconds and deploys on a free tier. "
         "A guiding principle is that every prediction reports its reasoning: nothing is an opaque black "
         "box. The models are trained on a seeded, reproducible synthetic history of roughly 5,800 "
         "deliveries spanning 180 days.")

    h2(doc, "7.1 ETA and Late-Risk Prediction")
    para(doc,
         "Three gradient-boosting models estimate operational outcomes for a shipment: the drop-off time "
         "(minutes to hand the parcel to the customer), the pickup time at the store, and the probability "
         "that a delivery misses its service-level agreement. A single shared feature pipeline "
         "(haversine distance, time-of-day and day-of-week traffic factor, store congestion, vehicle speed "
         "and the promised SLA) feeds both training and live serving, eliminating training/serving skew. "
         "Each prediction is accompanied by an exact additive explanation produced by a custom "
         "TreeContributionExplainer, which walks the gradient-boosting trees to recover the signed "
         "contribution of every feature.")

    h2(doc, "7.2 Validation: Cross-Validation and Baselines")
    para(doc,
         "To ensure the reported accuracy reflects genuine generalisation rather than a lucky train/test "
         "split, every model is validated with k-fold cross-validation and compared against a naive "
         "baseline. The regressors use 5-fold cross-validation; the late-risk classifier uses stratified "
         "5-fold; and the forecaster uses rolling-origin (time-series) cross-validation, which only ever "
         "predicts the future from the past. The tight cross-validation standard deviations demonstrate "
         "that the scores are stable across folds, and each model comfortably clears its baseline, "
         "confirming it has learned real signal.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(["Model", "Headline score", "Cross-validated", "Baseline", "Skill vs baseline"]):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    ml_rows = [
        ("Drop-off ETA (regressor)", "MAE 2.74 min, R2 0.94", "CV MAE 2.75 +/- 0.04", "Mean-predictor 9.78 min", "72% lower error"),
        ("Pick-up ETA (regressor)", "MAE 1.26 min, R2 0.63", "CV MAE 1.29 +/- 0.02", "Mean-predictor", "40% lower error"),
        ("Late-risk (classifier)", "ROC-AUC 0.82", "CV AUC 0.83 +/- 0.005", "Chance AUC 0.50", "Clears chance"),
        ("Demand forecast", "Orders MAPE 14.4%", "Rolling-origin CV 16.0 +/- 3.7%", "Seasonal-naive 18.5%", "Beats naive"),
    ]
    for r0, r1, r2, r3, r4 in ml_rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text, cells[4].text = r0, r1, r2, r3, r4
    para(doc, "")

    h2(doc, "7.3 Demand and Cost Forecasting")
    para(doc,
         "A deliberately white-box SeasonalTrendForecaster decomposes daily order volume and cost into a "
         "linear trend plus additive weekly seasonality, with a 95% prediction band from the residual "
         "spread. Because every component is explicit, the forecast explains itself in plain language "
         "(baseline level, growth per day, best and worst weekdays and the uncertainty band) rather than "
         "hiding behind an opaque model.")

    h2(doc, "7.4 Handling-Note NLP")
    para(doc,
         "Free-text delivery instructions are analysed by a hybrid multi-label classifier over a ten-tag "
         "taxonomy (fragile, do-not-stack, time-window, call-ahead and so on). A regex lexicon layer gives "
         "reliable coverage, and a learned TF-IDF plus one-vs-rest logistic-regression model generalises to "
         "unseen wording (for example \u2018the vase can shatter\u2019 maps to fragile). The model also "
         "extracts explicit delivery time windows and target floors, and each tag reports the exact token "
         "contributions that triggered it.")

    h2(doc, "7.5 Learning-to-Route Neural Policy")
    para(doc,
         "A pure-numpy pointer-network-style policy learns to sequence a courier\u2019s stops using a "
         "linear-softmax attention over seven scale-normalised geometric features. It is trained with "
         "REINFORCE and a greedy-rollout baseline using an exact softmax policy gradient (no autodiff). At "
         "inference, best-of-N sampled rollouts produce tours roughly 5% shorter than nearest-neighbour and "
         "within about 2% of a 2-opt reference, and every pick reports its top feature contributions. This "
         "demonstrates a modern reinforcement-learning approach to routing alongside the classical "
         "heuristics of Chapter 6.")

    h2(doc, "7.6 Courier-Behaviour Personas")
    para(doc,
         "Simulated GPS traces are mined for stop events (from dwell time and speed) and summarised into "
         "behavioural features, which K-Means then clusters into interpretable courier personas ranked by "
         "productivity. The clustering is evaluated with a silhouette score and, against the known "
         "simulation archetypes, a very high adjusted Rand index \u2014 an honest note in the report is that "
         "this validation is circular because the ground truth is synthetic.")

    h2(doc, "7.7 Feedback Loop, Drift and Chain-of-Custody Audit")
    para(doc,
         "Predictions are logged at dispatch and resolved at delivery, so the system measures live "
         "predicted-versus-actual error, week-over-week drift (the retraining trigger) and late-risk "
         "calibration. Independently, every status change appends a SHA-256 hash-chained handoff record "
         "(merchant to hub to courier to customer) that can be recomputed to detect tampering, and a GIS "
         "geofence confirms the delivery occurred at the destination.")


def engineering(doc):
    h1(doc, "8. Engineering, Security and Testing")
    h2(doc, "8.1 Security")
    bullets(doc, [
        "Passwords hashed with Werkzeug (PBKDF2); CSRF protection on every browser form.",
        "Role-based access control enforced by a decorator; server-side validation and open-redirect protection.",
    ])
    h2(doc, "8.2 Testing")
    para(doc,
         "An automated suite of 74 tests covers the full lifecycle (create \u2192 optimise \u2192 deliver "
         "\u2192 track), role protection, the traffic and closure model, dispatch-time planning, capacity "
         "limits, a dedicated test that verifies every optimisation technique \u2014 including the NetworkX "
         "strategies \u2014 returns a valid route that beats the FIFO baseline, and the entire predictive "
         "layer: the ETA/forecast models, the chain-of-custody and GIS audit, the handling-note NLP and "
         "assistant, the neural router and the courier-behaviour personas.")


def conclusion(doc):
    h1(doc, "9. Conclusion and Future Work")
    para(doc,
         "SwiftRoute delivers a complete logistics platform whose optimisation engine is not merely "
         "plausible but empirically validated. Benchmarking six techniques shows that automated planning "
         "cuts courier distance by around 60% versus naive dispatch and reaches within a fraction of a "
         "percent of optimal, while the technique catalogue spans a formal approximation algorithm and a "
         "metaheuristic for academic breadth. On top of this, an explainable predictive layer \u2014 "
         "cross-validated ETA and late-risk models, a transparent demand forecaster, handling-note NLP, a "
         "learning-to-route neural policy and courier-behaviour clustering \u2014 turns the platform into a "
         "genuine data-science system in which every prediction reports its reasoning.")
    para(doc, "Future work includes:")
    bullets(doc, [
        "A true capacitated VRP solver (e.g. Google OR-Tools) with time windows.",
        "Validating the predictive models on real operational data to complement the synthetic study.",
        "Live GPS tracking and dynamic re-routing around real-time incidents.",
        "Notifications (SMS/email) and cash-on-delivery settlement.",
    ])


def references(doc):
    h1(doc, "References")
    refs = [
        "Applegate, D. L., Bixby, R. E., Chv\u00e1tal, V., & Cook, W. J. (2006). The Traveling Salesman Problem: A Computational Study. Princeton University Press.",
        "Christofides, N. (1976). Worst-case analysis of a new heuristic for the travelling salesman problem. Technical Report, CMU.",
        "Lin, S., & Kernighan, B. W. (1973). An effective heuristic algorithm for the traveling-salesman problem. Operations Research, 21(2).",
        "Kirkpatrick, S., Gelatt, C. D., & Vecchi, M. P. (1983). Optimization by simulated annealing. Science, 220(4598).",
        "Hagberg, A., Schult, D., & Swart, P. (2008). Exploring network structure, dynamics, and function using NetworkX. Proceedings of SciPy.",
        "OSRM \u2014 Open Source Routing Machine. https://project-osrm.org",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Number")


# --------------------------------------------------------------------------- #
#  Build
# --------------------------------------------------------------------------- #
def build():
    summary = load_summary()
    doc = Document()

    # Base body font.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title_page(doc)
    abstract(doc)
    toc_page(doc)
    introduction(doc)
    background(doc)
    architecture(doc)
    methodology(doc)
    experiment(doc, summary)
    results(doc, summary)
    ml_layer(doc)
    engineering(doc)
    conclusion(doc)
    references(doc)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    figs = "with benchmark charts" if summary else "(no benchmark summary found \u2014 run the benchmark first)"
    print(f"Report written to {OUT} {figs}.")


if __name__ == "__main__":
    build()
